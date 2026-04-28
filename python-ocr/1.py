# Please make sure the requests library is installed
# pip install requests
import base64
import os
from pathlib import Path

import requests

try:
    from docx import Document

    DOCX_SUPPORT = True
except ImportError:
    Document = None
    DOCX_SUPPORT = False


API_URL = "https://95q8z9fflep806s7.aistudio-app.com/layout-parsing"
TOKEN = "9bc4642331c4a2bf69057eacd5f651d475813fcc"
FILE_PATH = Path(r"C:\Users\张海松\Desktop\关于印发《2025一2026年度大学生志愿服务西部计划实施方案》的通知.pdf")
OUTPUT_DIR = Path(__file__).with_name("output")


def encode_file_to_base64(file_path):
    with open(file_path, "rb") as file:
        return base64.b64encode(file.read()).decode("ascii")


def build_payload(file_data, file_type):
    return {
        "file": file_data,
        "fileType": file_type,
        "useDocOrientationClassify": False,
        "useDocUnwarping": False,
        "useChartRecognition": False,
    }


def request_layout_parsing(file_path, file_type=0):
    headers = {
        "Authorization": f"token {TOKEN}",
        "Content-Type": "application/json",
    }
    payload = build_payload(encode_file_to_base64(file_path), file_type)
    response = requests.post(API_URL, json=payload, headers=headers, timeout=120)
    response.raise_for_status()
    return response.json()["result"]


def save_page_markdown(page_index, markdown_text, output_dir):
    md_filename = output_dir / f"doc_{page_index}.md"
    md_filename.write_text(markdown_text, encoding="utf-8")
    print(f"Markdown document saved at {md_filename}")
    return md_filename


def download_markdown_images(images, output_dir):
    for img_path, img_url in images.items():
        full_img_path = output_dir / img_path
        full_img_path.parent.mkdir(parents=True, exist_ok=True)
        img_bytes = requests.get(img_url, timeout=120).content
        full_img_path.write_bytes(img_bytes)
        print(f"Image saved to: {full_img_path}")


def download_output_images(output_images, page_index, output_dir):
    for img_name, img_url in output_images.items():
        img_response = requests.get(img_url, timeout=120)
        if img_response.status_code == 200:
            filename = output_dir / f"{img_name}_{page_index}.jpg"
            filename.write_bytes(img_response.content)
            print(f"Image saved to: {filename}")
        else:
            print(f"Failed to download image, status code: {img_response.status_code}")


def merge_markdown_documents(page_markdowns, output_dir):
    merged_path = output_dir / "full.md"
    merged_text = "\n\n---\n\n".join(page_markdowns).strip()
    merged_path.write_text(f"{merged_text}\n", encoding="utf-8")
    print(f"Merged Markdown saved at {merged_path}")
    return merged_path


def export_docx_from_markdown(markdown_text, output_dir):
    if not DOCX_SUPPORT:
        print("python-docx is not installed; skipping DOCX export.")
        return None

    doc = Document()
    for block in markdown_text.split("\n\n"):
        stripped = block.strip()
        if not stripped:
            continue
        if stripped == "---":
            doc.add_page_break()
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        doc.add_paragraph(stripped)

    docx_path = output_dir / "full.docx"
    doc.save(docx_path)
    print(f"DOCX document saved at {docx_path}")
    return docx_path


def save_layout_parsing_results(result, output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    page_markdowns = []

    for i, res in enumerate(result["layoutParsingResults"]):
        markdown_text = res["markdown"]["text"]
        page_markdowns.append(markdown_text)
        save_page_markdown(i, markdown_text, output_dir)
        download_markdown_images(res["markdown"]["images"], output_dir)
        download_output_images(res["outputImages"], i, output_dir)

    merged_path = merge_markdown_documents(page_markdowns, output_dir)
    export_docx_from_markdown(merged_path.read_text(encoding="utf-8"), output_dir)
    return merged_path


def main():
    result = request_layout_parsing(FILE_PATH)
    save_layout_parsing_results(result, OUTPUT_DIR)


if __name__ == "__main__":
    main()
